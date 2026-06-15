from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.exports.shared import medicare_url, report_rows
from app.models import ReportPayload


def build_docx(payload: ReportPayload) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    _centered(document, "INFINITE — Managed by MEDELITE", 14, bold=True)
    _centered(document, "FACILITY ASSESSMENT SNAPSHOT", 12, bold=True)
    _centered(document, payload.facility.state or "", 11, bold=True)

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for label, value in report_rows(payload):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        _shade(cells[0], "F3F4F6")
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                paragraph.paragraph_format.space_after = Pt(0)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True

    document.add_paragraph()
    p = document.add_paragraph()
    _add_hyperlink(p, "Medicare Care Compare source profile", medicare_url(payload.facility.ccn))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _centered(document: DocxDocument, text: str, size: int, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(17, 24, 39)
    paragraph.paragraph_format.space_after = Pt(2)


def _shade(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1D4ED8")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(color)
    run_props.append(underline)
    run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
